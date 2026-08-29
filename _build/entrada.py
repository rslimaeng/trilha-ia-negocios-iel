# -*- coding: utf-8 -*-
"""Rito de entrada: leia a pasta do cliente e diga o que falta.

Rafael abre uma pasta com a ementa, os prints e o que mais tiver, e diz "lê essa
pasta e executa". Este passo é o que vem antes do "executa": ele inventaria o que
existe, confere contra o que o padrão precisa para começar, e **lista o que falta**.

🔴 Ele não decide nada e não escreve material. Ele só impede o começo cego, que é
escrever seis aulas em cima de uma ementa e descobrir na última que ninguém sabia
quem estava na sala.

Rodar:  python3 _build/entrada.py <pasta do cliente>
"""
import io
import os
import re
import subprocess
import sys

LEGIVEIS = {".md", ".txt", ".docx", ".pdf", ".xlsx", ".csv"}
IMAGENS = {".png", ".jpg", ".jpeg", ".webp", ".gif"}
IGNORAR = {".DS_Store"}


# ---------------------------------------------------------------------------
# O QUE O PADRÃO PRECISA PARA COMEÇAR
#
# Cada item traz os sinais que provam que ele existe na pasta. Sinal é indício,
# não prova: o relatório diz "achei indício de", nunca "está resolvido".
# ---------------------------------------------------------------------------
# Palavra inteira, sempre. "cor" casou dentro de "corte de conhecimento" e
# "manual" dentro de "manualmente": o relatório deu tudo verde numa pasta que não
# tinha a paleta do cliente. Check que sempre passa é check que não existe.
def cita(texto, termo):
    return re.search(r"(?<![\wÀ-ÿ])" + re.escape(termo) + r"(?![\wÀ-ÿ])", texto) is not None


# ---------------------------------------------------------------------------
# O QUE O PADRÃO PRECISA PARA COMEÇAR
#
# 🔴 Dois tipos, e a diferença é o que salva o relatório de mentir:
#
#   AUTO      vocabulário distintivo o bastante para a busca provar. Exige DOIS
#             sinais diferentes, porque um sozinho é coincidência.
#   CONFIRMAR nenhuma palavra prova. "marca" tanto é a paleta do cliente quanto
#             "formatação de marca" num guia de prompt. Estes SEMPRE viram
#             pergunta, mesmo com a pasta cheia.
# ---------------------------------------------------------------------------
PRECISA = [
    ("auto", "a ementa",
     ["ementa", "módulo", "modulo", "conteúdo programático", "programa",
      "objetivos do módulo", "trilha"],
     "sem ela não dá para recortar as aulas"),
    ("auto", "a carga horária e o formato",
     ["carga horária", "carga horaria", "encontro", "presencial", "híbrido",
      "hibrido", "mentoria", "turma", "40h", "aulas"],
     "decide quantas aulas cabem, e se tem exercício ao vivo"),
    ("auto", "o contexto do negócio",
     ["setor", "indústria", "industria", "processo", "rotina", "desafio",
      "área", "area"],
     "sem isso o caso e o insumo saem genéricos, e genérico não engaja"),
    ("auto", "um caso ou número real",
     ["roi", "kpi", "indicador", "meta", "planilha", "relatório", "relatorio"],
     "é o que faz o exercício valer. Piso: um número que a sala reconheça"),

    ("confirmar", "quem está na sala",
     [], "é a ORIGEM do GPS. Cargo, o que já usa de IA, e o que NÃO sabe"),
    ("confirmar", "a marca do cliente",
     [], "vira o marca.css. Peça o hex da cor principal, não o logo"),
    ("confirmar", "o que o cliente já tentou e não deu certo",
     [], "alimenta o .diagnostico e desarma a objeção antes do exercício"),
    ("confirmar", "quem aprova o material, e até quando",
     [], "material aprovado tarde vira material refeito"),
]


def texto_de(caminho):
    ext = os.path.splitext(caminho)[1].lower()
    try:
        if ext in (".md", ".txt", ".csv"):
            return io.open(caminho, encoding="utf-8", errors="ignore").read()
        if ext == ".docx":
            import docx
            d = docx.Document(caminho)
            partes = [p.text for p in d.paragraphs]
            for t in d.tables:
                for linha in t.rows:
                    partes.extend(c.text for c in linha.cells)
            return "\n".join(partes)
        if ext == ".pdf":
            r = subprocess.run(["pdftotext", "-q", caminho, "-"],
                               capture_output=True, timeout=60)
            return r.stdout.decode("utf-8", "ignore")
        if ext == ".xlsx":
            import openpyxl
            wb = openpyxl.load_workbook(caminho, read_only=True, data_only=True)
            partes = []
            for ws in wb.worksheets:
                partes.append(ws.title)
                for i, linha in enumerate(ws.iter_rows(values_only=True)):
                    if i > 40:
                        break
                    partes.extend(str(c) for c in linha if c is not None)
            return "\n".join(partes)
    except Exception as e:
        return "\x00ERRO\x00" + str(e)
    return ""


def versao(nome):
    """v3 ganha de v1. Ler a versão velha e produzir em cima dela é um jeito
    silencioso de entregar o que o cliente já descartou."""
    m = re.search(r"[_\-\s]v(\d+)", nome, re.I)
    return int(m.group(1)) if m else None


def main():
    if len(sys.argv) < 2:
        sys.exit("uso: python3 _build/entrada.py <pasta do cliente>")
    pasta = os.path.abspath(sys.argv[1])
    if not os.path.isdir(pasta):
        sys.exit("não é uma pasta: " + pasta)

    arquivos, imagens, ilegiveis = [], [], []
    for raiz, _, nomes in os.walk(pasta):
        for n in sorted(nomes):
            if n in IGNORAR or n.startswith("."):
                continue
            caminho = os.path.join(raiz, n)
            ext = os.path.splitext(n)[1].lower()
            rel = os.path.relpath(caminho, pasta)
            if ext in IMAGENS:
                imagens.append((rel, os.path.getsize(caminho)))
            elif ext in LEGIVEIS:
                arquivos.append((rel, caminho, ext))
            else:
                ilegiveis.append(rel)

    print("PASTA · %s\n" % pasta)
    print("%d arquivo(s) de texto · %d imagem(ns) · %d que eu não leio\n"
          % (len(arquivos), len(imagens), len(ilegiveis)))

    # ---- o que dá para ler
    corpus, familias = [], {}
    print("O QUE EU CONSIGO LER")
    for rel, caminho, ext in arquivos:
        t = texto_de(caminho)
        if t.startswith("\x00ERRO\x00"):
            print("  FALHOU  %-46s %s" % (rel, t[7:][:60]))
            continue
        palavras = len(t.split())
        corpus.append(t.lower())
        print("  ok      %-46s %6d palavras" % (rel, palavras))
        base = re.sub(r"[_\-\s]v\d+", "", os.path.splitext(rel)[0], flags=re.I)
        familias.setdefault(base, []).append((versao(rel), rel))
    print()

    # ---- versões
    duplas = {b: v for b, v in familias.items() if len(v) > 1}
    if duplas:
        print("VERSÕES DO MESMO DOCUMENTO")
        for base, vs in duplas.items():
            vs = sorted(vs, key=lambda x: (x[0] is None, x[0]))
            mais_nova = vs[-1][1]
            print("  %s" % base)
            for n, rel in vs:
                print("     %s %s" % ("→ USE ESTA:" if rel == mais_nova
                                      else "   ignore: ", rel))
        print("  🔴 Produzir em cima da versão velha entrega o que o cliente já"
              " descartou.\n")

    # ---- imagens
    if imagens:
        print("IMAGENS · eu preciso ABRIR uma a uma, não dá para varrer")
        for rel, tam in imagens:
            print("  %-56s %6d KB" % (rel, tam // 1024))
        print("  Print costuma ser página de vendas, print de ferramenta ou"
              " referência visual.\n")

    if ilegiveis:
        print("NÃO CONSIGO LER (peça conversão ou o conteúdo em texto)")
        for rel in ilegiveis:
            print("  " + rel)
        print()

    # ---- o que falta
    tudo = "\n".join(corpus)
    print("O QUE O PADRÃO PRECISA")
    faltando = []
    for tipo, nome, sinais, porque in PRECISA:
        if tipo == "confirmar":
            print("  PERGUNTE %-32s %s" % (nome, porque))
            faltando.append(nome)
            continue
        achou = [t for t in sinais if cita(tudo, t)]
        # um sinal sozinho é coincidência; dois já é vocabulário do assunto
        if len(achou) >= 2:
            print("  indício  %-32s (%s)" % (nome, ", ".join(achou[:3])))
        else:
            print("  FALTA    %-32s %s" % (nome, porque))
            faltando.append(nome)
    print()

    if faltando:
        print("PERGUNTE AO RAFAEL, ANTES DE ESCREVER")
        for i, f in enumerate(faltando[:3], 1):
            print("  %d. %s" % (i, f))
        if len(faltando) > 3:
            print("\n  🔴 Faltam %d itens, e o teto por turno é 3. Briefing fraco:"
                  "\n     peça um norte melhor em vez de bombardear com pergunta."
                  % len(faltando))
            print("     Os outros: %s" % ", ".join(faltando[3:]))
    else:
        print("Os itens automáticos têm indício. Os de confirmação continuam"
              " valendo.")

    print("\n" + "-" * 68)
    print("Indício NÃO é prova. Abra os arquivos e confirme antes de produzir,")
    print("e aplique a skill leitura-de-fonte no material que for ementa ou")
    print("transcrição. Depois siga o COMO-EXECUTAR.md a partir do passo 1.")


if __name__ == "__main__":
    main()
