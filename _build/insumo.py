# -*- coding: utf-8 -*-
"""Gerador do insumo: o arquivo que o aluno baixa.

O `.xlsx` e o PROMPT da página de caso estão acoplados: o prompt cita as colunas
do arquivo. Quando os dois nascem à mão, em momentos diferentes, eles divergem, e
o defeito não aparece em lugar nenhum: o aluno cola o prompt, a IA responde sobre
uma coluna que não existe, e a resposta sai plausível e errada na frente da sala.

Aqui o arquivo nasce de uma especificação declarada, e o `insumos.json` que sai
junto é o que o gate confere contra a página.

🔴 A sujeira é DECLARADA. Cada armadilha existe para provocar um erro específico,
e toda armadilha aplicada tem de aparecer em algum lugar do material: no prompt,
no gabarito ou nas pegadinhas. Armadilha que ninguém explica depois não é
pedagogia, é pegadinha de prova.

🔴 Tudo é determinístico. Sem semente fixa, cada execução geraria um arquivo
diferente, a turma receberia planilhas que não batem entre si, e o gabarito
deixaria de valer.

Rodar:  python3 _build/insumo.py
"""
import io
import json
import os
import random
import re
import datetime

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter

AQUI = os.path.dirname(os.path.abspath(__file__))
RAIZ = os.path.dirname(AQUI)
DESTINO = os.path.join(RAIZ, "_arquivos")

# A semente é o que garante que a planilha de hoje é a planilha de amanhã.
SEMENTE = 20260820


# ---------------------------------------------------------------------------
# O CATÁLOGO DE ARMADILHAS
#
# Cada uma diz o que ensina. Esse texto não é comentário: ele sai no relatório
# para quem vai escrever a aula, porque é ele que precisa reaparecer no material.
# ---------------------------------------------------------------------------
ARMADILHAS = {
    "linha_em_branco_antes_do_cabecalho":
        "que o arquivo precisa ser olhado antes de ser usado",
    "data_em_dois_formatos":
        "que a IA ordena errado sem avisar",
    "nome_em_dois_jeitos":
        "que juntar por semelhança é decisão, não detalhe. É a origem do "
        "parágrafo 'Na dúvida' do prompt",
    "coluna_com_espaco_no_nome":
        "que o nome que ela vê não é o nome que a ferramenta lê",
    "numero_como_texto":
        "que soma que não soma tem causa, e a causa é achável",
    "linha_duplicada":
        "que conferir o total contra a soma é barato",
    "valor_faltando":
        "que o buraco tem de ser declarado, nunca estimado em silêncio",
}


# ---------------------------------------------------------------------------
# O CATÁLOGO DE ARMADILHAS DE DOCUMENTO
#
# 🔴 São outras, e é por isso que existem duas listas. Planilha erra em COLUNA:
# formato, tipo, duplicata. Documento erra em SENTIDO: contradiz a si mesmo,
# omite o dono, e diz "próxima semana" sem dizer a partir de quando.
#
# Um exercício de documento com armadilha de planilha ensina a coisa errada: a
# pessoa sai procurando célula quando o problema é o parágrafo.
# ---------------------------------------------------------------------------
ARMADILHAS_DOC = {
    "decisao_registrada_duas_vezes":
        "que documento longo se contradiz, e que a IA repete a primeira "
        "versão que encontra sem avisar que há uma segunda",
    "responsavel_sem_nome":
        "que pendência sem dono não é pendência, é intenção",
    "prazo_relativo":
        "que 'na próxima semana' perde o sentido no dia em que alguém lê o "
        "documento fora da semana em que ele foi escrito",
    "sigla_nunca_expandida":
        "que a IA preenche a sigla que ela não conhece com a que ela conhece",
    "numero_que_nao_bate":
        "que o total escrito na frase e a soma dos itens da lista são duas "
        "afirmações diferentes, e só uma delas foi conferida",
    "paragrafo_de_outra_reuniao":
        "que copiar e colar traz o contexto errado junto, e ninguém relê",
}


# ---------------------------------------------------------------------------
# A ESPECIFICAÇÃO
#
# Trocar o insumo de um curso é mexer aqui. Mais nada.
# ---------------------------------------------------------------------------
INSUMOS = {
    "fechamento-semanal-exemplo": dict(
        formato="xlsx",
        caso="caso",                       # a página que oferece o download
        titulo="Fechamento semanal · 12 lojas",
        semanas=4,
        lojas=["Centro", "Norte", "Sul", "Litoral", "Shopping", "Rodoviária",
               "Bairro Alto", "Praia", "Industrial", "Universidade",
               "Terminal", "Feira"],
        produtos=[("Sorvete 2L", 34.90), ("Picolé unidade", 6.50),
                  ("Açaí 500ml", 22.00), ("Sorvete 1L", 19.90),
                  ("Torta gelada", 48.00), ("Milkshake", 18.50),
                  ("Casquinha", 9.00), ("Pote família 5L", 79.90)],
        armadilhas=list(ARMADILHAS),
        aviso=("Os dados deste arquivo são FICTÍCIOS. As lojas, os valores e os "
               "produtos foram inventados para o treinamento. Nenhum dado de "
               "cliente real foi usado."),
    ),
    "ata-do-projeto-exemplo": dict(
        formato="docx",
        caso="caso",
        titulo="Ata · Comitê do projeto Guarita",
        armadilhas=list(ARMADILHAS_DOC),
        aviso=("Este documento é FICTÍCIO. A empresa, as pessoas, os números e "
               "as decisões foram inventados para o treinamento. Nenhum "
               "documento de cliente real foi usado."),
    ),
}


# ---------------------------------------------------------------------------
# OS DADOS
# ---------------------------------------------------------------------------
def monta_vendas(spec, r):
    """Uma linha por venda. É o volume que faz a demonstração ter argumento:
    numa planilha de 15 linhas alguém pensa 'isso eu fazia na mão'."""
    inicio = datetime.date(2026, 7, 6)
    linhas = []
    cupom = 0
    for semana in range(spec["semanas"]):
        for dia in range(7):
            data = inicio + datetime.timedelta(days=semana * 7 + dia)
            for loja in spec["lojas"]:
                # fim de semana vende mais; é o que faz a leitura ter o que achar
                n = r.randint(9, 14) if data.weekday() >= 5 else r.randint(5, 9)
                for _ in range(n):
                    produto, preco = r.choice(spec["produtos"])
                    qtd = r.randint(1, 6)
                    cupom += 1
                    linhas.append({
                        "cupom": "C%06d" % cupom,
                        "data": data,
                        "loja": loja,
                        "produto": produto,
                        "qtd": qtd,
                        "unitario": preco,
                        "total": round(qtd * preco, 2),
                        "pagamento": r.choice(["pix", "credito", "debito", "dinheiro"]),
                    })
    return linhas


def suja(linhas, spec, r):
    """Aplica as armadilhas declaradas. Devolve (linhas, o que foi aplicado)."""
    aplicadas = []
    quais = set(spec["armadilhas"])

    if "nome_em_dois_jeitos" in quais:
        # a MESMA loja escrita de três jeitos, em 18% das linhas
        for l in linhas:
            if l["loja"] == "Centro" and r.random() < 0.18:
                l["loja"] = r.choice(["centro", "CENTRO", "Loja Centro"])
        aplicadas.append("nome_em_dois_jeitos")

    if "data_em_dois_formatos" in quais:
        # um terço vira texto dd/mm/aaaa; o resto fica como data de verdade
        for l in linhas:
            if r.random() < 0.33:
                l["data"] = l["data"].strftime("%d/%m/%Y")
        aplicadas.append("data_em_dois_formatos")

    if "numero_como_texto" in quais:
        for l in linhas:
            if r.random() < 0.12:
                l["total"] = ("%.2f" % l["total"]).replace(".", ",")
        aplicadas.append("numero_como_texto")

    if "valor_faltando" in quais:
        for l in linhas:
            if r.random() < 0.02:
                l["qtd"] = None
        aplicadas.append("valor_faltando")

    if "linha_duplicada" in quais:
        # UMA linha, no meio, com o mesmo número de cupom: é assim que um export
        # duplicado se parece, e é o que torna a duplicata achável de verdade
        i = len(linhas) // 3
        linhas.insert(i + 1, dict(linhas[i]))
        aplicadas.append("linha_duplicada")

    return linhas, aplicadas


# ---------------------------------------------------------------------------
# A PLANILHA
# ---------------------------------------------------------------------------
CABECALHO = PatternFill("solid", fgColor="DDDDDD")


def grava_xlsx(slug, spec):
    r = random.Random(SEMENTE)
    linhas = monta_vendas(spec, r)
    linhas, aplicadas = suja(linhas, spec, r)
    quais = set(spec["armadilhas"])

    wb = Workbook()
    # 🔴 openpyxl carimba datetime.now() no docProps, e o arquivo passa a mudar a
    # cada execução mesmo com os dados idênticos. Quem confere por hash conclui
    # que a planilha mudou, e a turma acaba com versões que ninguém consegue
    # comparar. O carimbo é fixo, e é a data de referência do próprio insumo.
    wb.properties.creator = "Gerador de insumo · padrão de treinamento"
    wb.properties.created = datetime.datetime(2026, 8, 20, 12, 0, 0)
    wb.properties.modified = wb.properties.created

    # ---- aba 1 · leia-me. Vem primeiro porque é a que abre.
    ws = wb.active
    ws.title = "leia-me"
    ws["A1"] = spec["titulo"]
    ws["A1"].font = Font(bold=True, size=14)
    ws["A3"] = spec["aviso"]
    ws["A3"].alignment = Alignment(wrap_text=True, vertical="top")
    ws.merge_cells("A3:F6")
    ws["A8"] = "O que tem em cada aba"
    ws["A8"].font = Font(bold=True)
    guia = [("vendas", "uma linha por venda, item a item"),
            ("resumo", "o fechamento semana a semana"),
            ("produtos", "o cadastro, com o preço de tabela")]
    for i, (aba, o_que) in enumerate(guia, start=9):
        ws.cell(row=i, column=1, value=aba).font = Font(bold=True)
        ws.cell(row=i, column=2, value=o_que)
    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 52

    # ---- aba 2 · vendas, com as armadilhas
    ws = wb.create_sheet("vendas")
    # a linha em branco antes do cabeçalho: o arquivo precisa ser olhado
    primeira = 3 if "linha_em_branco_antes_do_cabecalho" in quais else 1
    if primeira == 3:
        ws["A1"] = "Relatório exportado do sistema · não alterar"
        ws["A1"].font = Font(italic=True, color="888888")

    # a coluna com espaço no nome: o nome que ela vê não é o que a ferramenta lê
    col_total = " Valor Total" if "coluna_com_espaco_no_nome" in quais else "Valor Total"
    colunas = ["Cupom", "Data", "Loja", "Produto", "Qtd", "Valor Unitario",
               col_total, "Forma de Pagamento"]
    for c, nome in enumerate(colunas, start=1):
        cel = ws.cell(row=primeira, column=c, value=nome)
        cel.font = Font(bold=True)
        cel.fill = CABECALHO
    for i, l in enumerate(linhas, start=primeira + 1):
        ws.cell(row=i, column=1, value=l["cupom"])
        ws.cell(row=i, column=2, value=l["data"])
        ws.cell(row=i, column=3, value=l["loja"])
        ws.cell(row=i, column=4, value=l["produto"])
        ws.cell(row=i, column=5, value=l["qtd"])
        ws.cell(row=i, column=6, value=l["unitario"])
        ws.cell(row=i, column=7, value=l["total"])
        ws.cell(row=i, column=8, value=l["pagamento"])
    for c, larg in enumerate([11, 13, 15, 18, 7, 15, 13, 19], start=1):
        ws.column_dimensions[get_column_letter(c)].width = larg

    # ---- aba 3 · resumo, e ele NÃO bate com a soma da aba vendas de propósito:
    # é a linha duplicada aparecendo. Quem confere acha; quem não confere, não.
    ws = wb.create_sheet("resumo")
    cab = ["Semana", "Periodo", "Faturamento", "Cupons", "Ticket medio",
           "Observacao da operacao"]
    for c, nome in enumerate(cab, start=1):
        cel = ws.cell(row=1, column=c, value=nome)
        cel.font = Font(bold=True)
        cel.fill = CABECALHO
    obs = ["semana de referencia", "produto novo chegou dia 15",
           "sabado com recorde de fluxo", "duas lojas sem produto na quarta"]
    inicio = datetime.date(2026, 7, 6)
    for s in range(spec["semanas"]):
        d0 = inicio + datetime.timedelta(days=s * 7)
        d1 = d0 + datetime.timedelta(days=6)
        # só as linhas com total numérico entram: as que viraram texto ficam de
        # fora, e é essa a diferença que o aluno vai caçar
        do_periodo = [l for l in linhas
                      if _na_semana(l["data"], d0, d1)
                      and isinstance(l["total"], (int, float))]
        fat = round(sum(l["total"] for l in do_periodo), 2)
        cupons = len(do_periodo)
        ws.cell(row=s + 2, column=1, value="S%d" % (s + 1))
        ws.cell(row=s + 2, column=2, value="%s a %s" % (d0.strftime("%d/%m"),
                                                        d1.strftime("%d/%m")))
        ws.cell(row=s + 2, column=3, value=fat)
        ws.cell(row=s + 2, column=4, value=cupons)
        ws.cell(row=s + 2, column=5, value=round(fat / cupons, 2) if cupons else 0)
        ws.cell(row=s + 2, column=6, value=obs[s])
    for c, larg in enumerate([10, 18, 16, 10, 14, 34], start=1):
        ws.column_dimensions[get_column_letter(c)].width = larg

    # ---- aba 4 · produtos
    ws = wb.create_sheet("produtos")
    for c, nome in enumerate(["Produto", "Preco de tabela", "Categoria"], start=1):
        cel = ws.cell(row=1, column=c, value=nome)
        cel.font = Font(bold=True)
        cel.fill = CABECALHO
    for i, (nome, preco) in enumerate(spec["produtos"], start=2):
        ws.cell(row=i, column=1, value=nome)
        ws.cell(row=i, column=2, value=preco)
        ws.cell(row=i, column=3, value="gelados")
    for c, larg in enumerate([22, 18, 14], start=1):
        ws.column_dimensions[get_column_letter(c)].width = larg

    # as duas armadilhas que moram na gravação, e não no tratamento dos dados.
    # Elas ficavam de fora da lista, e o manifesto saía mentindo por omissão.
    if primeira == 3:
        aplicadas.append("linha_em_branco_antes_do_cabecalho")
    if col_total.startswith(" "):
        aplicadas.append("coluna_com_espaco_no_nome")

    os.makedirs(DESTINO, exist_ok=True)
    caminho = os.path.join(DESTINO, slug + ".xlsx")
    wb.save(caminho)
    _congela_o_zip(caminho)

    return {
        "arquivo": slug + ".xlsx",
        "formato": "xlsx",
        "caso": spec["caso"],
        "titulo": spec["titulo"],
        "linhas": len(linhas),
        "abas": [w.title for w in wb.worksheets],
        "colunas": {"vendas": colunas, "resumo": cab,
                    "produtos": ["Produto", "Preco de tabela", "Categoria"]},
        "armadilhas": {a: ARMADILHAS[a] for a in aplicadas},
        "cabecalho_na_linha": primeira,
    }


# ---------------------------------------------------------------------------
# O DOCUMENTO
#
# 🔴 O texto é FIXO, não sorteado. Numa planilha a aleatoriedade com semente é
# o que dá volume sem inventar caso a caso; num documento ela produziria frases
# que ninguém leu. O documento é escrito à mão aqui, e é conferido aqui: a
# função que grava também MEDE a armadilha que aplicou, e o que ela não
# conseguir medir não entra no manifesto.
#
# 🔴 Nenhum nome, empresa ou número deste documento é real. A regra de
# segurança do padrão vale igual no insumo: nome de outro cliente não entra,
# nem anonimizado.
# ---------------------------------------------------------------------------

# O par que se contradiz. A primeira data está no corpo, a segunda no anexo --
# 40 linhas depois, que é onde a releitura desiste.
VIRADA_NO_CORPO = "15/09/2026"
VIRADA_NO_ANEXO = "29/09/2026"

# O total ESCRITO na frase, e os itens que a lista traz. Os dois números
# existem no documento; um deles está errado, e o documento não diz qual.
TOTAL_ESCRITO = 148000
ITENS_DO_ORCAMENTO = [
    ("Licenças e assinaturas", 46000),
    ("Consultoria de implantação", 72000),
    ("Treinamento das equipes", 28000),
    ("Reserva técnica", 16000),
]

CORPO_DA_ATA = [
    ("h1", "Ata · Comitê do projeto Guarita"),
    ("meta", "Reunião ordinária · 26/08/2026 · 14h às 15h30 · sala virtual"),
    ("meta", "Participantes: Direção de Operações, Coordenação de TI, "
             "Coordenação Comercial, Escritório de Projetos"),

    ("h2", "1. Contexto"),
    ("p", "O comitê se reuniu para revisar o andamento do projeto Guarita, que "
          "substitui o controle de ocorrências feito hoje em planilha "
          "compartilhada. O piloto rodou em duas unidades por seis semanas."),
    ("p", "O PCO apresentou o consolidado do piloto. Segundo o PCO, o volume de "
          "ocorrências registradas subiu 34% em relação ao mesmo período do ano "
          "anterior, o que era esperado: parte do aumento é registro que antes "
          "não acontecia. O PCO recomenda manter a apuração separada por "
          "unidade até a virada."),

    ("h2", "2. Decisões"),
    ("p", "Ficou decidido que a virada das demais unidades acontece em "
          + VIRADA_NO_CORPO + ", condicionada ao aceite do relatório de piloto."),
    ("p", "Ficou decidido que o treinamento das equipes ocorre antes da virada, "
          "em turmas por unidade, e que a Coordenação de TI publica o "
          "calendário na próxima semana."),
    ("p", "Ficou decidido que o modelo de relatório mensal passa a ser o do "
          "PCO, e que o modelo antigo sai de circulação assim que o novo "
          "estiver publicado."),

    ("h2", "3. Pendências"),
    ("p", "A equipe vai avaliar se o campo de reincidência entra já na virada "
          "ou fica para a fase 2."),
    ("p", "Alguém do time comercial precisa confirmar se o contrato atual "
          "cobre as duas unidades novas."),
    ("p", "A Coordenação de TI fecha a lista de perfis de acesso até o fim do "
          "mês."),
    ("p", "Falta definir quem responde pelo dado quando a unidade não registra "
          "a ocorrência no prazo."),

    ("h2", "4. Orçamento aprovado"),
    ("p", "Os itens aprovados para o ciclo totalizam R$ " +
          "{:,.0f}".format(TOTAL_ESCRITO).replace(",", ".") + ", conforme a "
          "lista abaixo."),
    ("itens", None),
    ("p", "O desembolso segue o cronograma financeiro já acordado, sem "
          "antecipação."),

    ("h2", "5. Registro do encaminhamento anterior"),
    ("p", "Sobre o projeto Farol: a integração com o sistema de portaria "
          "continua parada aguardando o retorno do fornecedor, e o comitê "
          "decidiu não escalar antes do fim do contrato vigente. O "
          "acompanhamento segue com o Escritório de Projetos, em reunião "
          "quinzenal."),

    ("h2", "6. Anexo · extrato da ata anterior"),
    ("p", "Extrato mantido para referência, conforme praxe do comitê."),
    ("p", "Ficou decidido que a virada das demais unidades acontece em "
          + VIRADA_NO_ANEXO + ", após o encerramento do ciclo de férias da "
          "equipe de TI."),
    ("p", "O PCO fica responsável por consolidar os números do piloto e "
          "apresentar no próximo comitê."),
]


def grava_docx(slug, spec):
    """Grava o .docx e MEDE cada armadilha no documento gravado.

    🔴 Medir depois de escrever, e não confiar no que a spec declarou, é a mesma
    regra do .xlsx: a lista de armadilhas do manifesto é o que o arquivo TEM,
    não o que alguém pretendia pôr nele.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    secoes = []
    paragrafos = 0

    for tipo, texto in CORPO_DA_ATA:
        if tipo == "h1":
            doc.add_heading(texto, level=1)
            secoes.append(texto)
        elif tipo == "h2":
            doc.add_heading(texto, level=2)
            secoes.append(texto)
        elif tipo == "meta":
            p = doc.add_paragraph(texto)
            p.runs[0].font.size = Pt(10)
            paragrafos += 1
        elif tipo == "itens":
            for nome, valor in ITENS_DO_ORCAMENTO:
                doc.add_paragraph(
                    "{} — R$ {:,.0f}".format(nome, valor).replace(",", "."),
                    style="List Bullet")
                paragrafos += 1
        else:
            doc.add_paragraph(texto)
            paragrafos += 1

    doc.add_paragraph("")
    aviso = doc.add_paragraph(spec["aviso"])
    aviso.runs[0].font.size = Pt(9)
    paragrafos += 1

    # ---- a medição: o que o documento gravado realmente tem ----
    todo = "\n".join(p.text for p in doc.paragraphs)
    aplicadas = []
    if VIRADA_NO_CORPO in todo and VIRADA_NO_ANEXO in todo:
        aplicadas.append("decisao_registrada_duas_vezes")
    if "A equipe vai avaliar" in todo and "Alguém do time comercial" in todo:
        aplicadas.append("responsavel_sem_nome")
    if "na próxima semana" in todo and "até o fim do mês" in todo:
        aplicadas.append("prazo_relativo")
    # a sigla existe e NUNCA aparece expandida
    if todo.count("PCO") >= 4 and "Planejamento e Controle" not in todo:
        aplicadas.append("sigla_nunca_expandida")
    soma = sum(v for _, v in ITENS_DO_ORCAMENTO)
    if soma != TOTAL_ESCRITO:
        aplicadas.append("numero_que_nao_bate")
    if "projeto Farol" in todo:
        aplicadas.append("paragrafo_de_outra_reuniao")

    os.makedirs(DESTINO, exist_ok=True)
    caminho = os.path.join(DESTINO, slug + ".docx")
    doc.save(caminho)
    _congela_o_zip(caminho)

    return {
        "arquivo": slug + ".docx",
        "formato": "docx",
        "caso": spec["caso"],
        "titulo": spec["titulo"],
        "paragrafos": paragrafos,
        "secoes": secoes,
        "armadilhas": {a: ARMADILHAS_DOC[a] for a in aplicadas},
        "total_escrito": TOTAL_ESCRITO,
        "soma_dos_itens": soma,
    }


# 🔴 O .xlsx é um zip, e o zip guarda a hora de gravação DE CADA ENTRADA. O
# openpyxl usa a hora atual, então o arquivo trocava de hash a cada execução
# mesmo com todo o conteúdo idêntico.
#
# Eu afirmei duas vezes que o gerador era idempotente, e as duas verificações
# rodaram dentro do mesmo segundo. A terceira cruzou a virada e reprovou.
#
# Por que importa: sem hash estável, "a planilha mudou?" não tem resposta, o git
# marca o binário como modificado a cada rodada, e a turma não tem como conferir
# que está com o mesmo arquivo.
CARIMBO_DO_ZIP = (2026, 8, 20, 12, 0, 0)


def _congela_o_zip(caminho):
    """Regrava o zip com a data de cada entrada fixa, na mesma ordem."""
    import zipfile
    with zipfile.ZipFile(caminho) as z:
        itens = [(i, z.read(i.filename)) for i in z.infolist()]
    # 🔴 O openpyxl SOBRESCREVE dcterms:modified na hora de salvar, ignorando o
    # que foi definido em wb.properties. Só dá para fixar depois.
    quando = "%04d-%02d-%02dT%02d:%02d:%02dZ" % CARIMBO_DO_ZIP
    # 🔴 \1 seguido de "2026" vira \120, que o regex lê como escape octal e
    # produz o byte 'P'. O XML sai quebrado e o .xlsx para de abrir. Sempre
    # \g<1> quando o texto que vem depois começa com dígito.
    itens = [(i, re.sub(rb"(<dcterms:modified[^>]*>)[^<]*(</dcterms:modified>)",
                        rb"\g<1>" + quando.encode() + rb"\g<2>", d)
              if i.filename == "docProps/core.xml" else d)
             for i, d in itens]
    with zipfile.ZipFile(caminho, "w", zipfile.ZIP_DEFLATED) as z:
        for info, dados in itens:
            novo = zipfile.ZipInfo(info.filename, date_time=CARIMBO_DO_ZIP)
            novo.compress_type = info.compress_type
            novo.external_attr = info.external_attr
            z.writestr(novo, dados)


def _na_semana(data, d0, d1):
    if isinstance(data, str):
        data = datetime.datetime.strptime(data, "%d/%m/%Y").date()
    return d0 <= data <= d1


# Cada formato tem o seu gravador. Um curso pede planilha, outro pede documento,
# e há caso que pede OS DOIS -- a ata que decide e a planilha que prova. Nesse
# caso são duas entradas em INSUMOS apontando para o mesmo "caso", e não um
# terceiro formato: o gerador não precisa saber que elas andam juntas, a página
# de caso é que oferece as duas.
GRAVADORES = {
    "xlsx": grava_xlsx,
    "docx": grava_docx,
}


def main():
    manifesto = {}
    for slug, spec in INSUMOS.items():
        grava = GRAVADORES.get(spec["formato"])
        if not grava:
            print("  pulou:   %s (formato %s ainda não implementado)"
                  % (slug, spec["formato"]))
            continue
        info = grava(slug, spec)
        manifesto[slug] = info
        caminho = os.path.join(DESTINO, info["arquivo"])
        print("  gravado: %-34s %6d bytes" % (info["arquivo"],
                                              os.path.getsize(caminho)))
        if info["formato"] == "xlsx":
            print("           %d linhas · %d abas · cabeçalho na linha %d"
                  % (info["linhas"], len(info["abas"]),
                     info["cabecalho_na_linha"]))
        else:
            print("           %d parágrafos · %d seções · total escrito R$ %s "
                  "contra soma R$ %s"
                  % (info["paragrafos"], len(info["secoes"]),
                     "{:,.0f}".format(info["total_escrito"]).replace(",", "."),
                     "{:,.0f}".format(info["soma_dos_itens"]).replace(",", ".")))
        print("           armadilhas:")
        for a, ensina in info["armadilhas"].items():
            print("             · %-36s ensina %s" % (a, ensina))

        faltou = set(spec["armadilhas"]) - set(info["armadilhas"])
        if faltou:
            print("           🔴 declaradas e NÃO aplicadas: %s" % sorted(faltou))

    with io.open(os.path.join(AQUI, "insumos.json"), "w", encoding="utf-8") as f:
        f.write(json.dumps(manifesto, indent=2, ensure_ascii=False, sort_keys=True))
    print("\n  %d insumo(s) · manifesto em _build/insumos.json" % len(manifesto))
    print("\n  🔴 Toda armadilha da lista acima precisa aparecer no material:")
    print("     no prompt, no gabarito ou nas pegadinhas. O G32 confere o resto.")


if __name__ == "__main__":
    main()
